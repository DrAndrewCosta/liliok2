
import streamlit as st
import pandas as pd
from glob import glob
import os
from docx import Document
from datetime import date
from streamlit_js_eval import streamlit_js_eval
from lili_voz_comandos import LiliComandosVoz

st.set_page_config(page_title="Painel Lili", layout="wide")

if "texto_laudo" not in st.session_state:
    st.session_state.texto_laudo = ""

df_base = pd.read_csv("frases.csv", sep="|") if os.path.exists("frases.csv") else pd.DataFrame(columns=["nome_da_alteracao", "frase"])

col1, col2 = st.columns([1, 2])

with col1:
    st.header("📋 Dados do exame")

    templates_disponiveis = sorted([
        os.path.basename(path) for path in glob("templates/*.docx")
    ])
    template_escolhido = st.selectbox("🗂️ Template:", templates_disponiveis)

    nome_paciente = st.text_input("Nome do paciente", value=st.session_state.get("nome_paciente", ""))
    dn_paciente = st.date_input("Data de nascimento", value=st.session_state.get("dn_paciente", date.today()))
    medico_solicitante = st.text_input("Médico solicitante", value=st.session_state.get("medico_solicitante", ""))
    data_exame = st.date_input("Data do exame", value=st.session_state.get("data_exame", date.today()))

    st.session_state.update({
        "nome_paciente": nome_paciente,
        "dn_paciente": dn_paciente,
        "medico_solicitante": medico_solicitante,
        "data_exame": data_exame
    })

    fala = st.text_input("💬 Comando por voz (comece com 'Lili,'):")

    if fala and fala.lower().startswith("lili"):
        lili = LiliComandosVoz(df_base, responder_em_voz=True)
        novo_laudo, mensagem, acao = lili.processar_comando(fala, st.session_state.texto_laudo)
        if mensagem:
            st.success(mensagem)
        st.session_state.texto_laudo = novo_laudo

    csv_file = st.file_uploader("📤 Importar frases clínicas (.csv)", type=["csv"])
    if csv_file:
        try:
            df_csv = pd.read_csv(csv_file)
            df_novo = pd.concat([df_base, df_csv]).drop_duplicates(subset="frase", keep="first").reset_index(drop=True)
            df_novo.to_csv("frases.csv", index=False)
            st.success(f"{len(df_csv)} frases importadas com sucesso.")
        except Exception as e:
            st.error(f"Erro ao processar CSV: {e}")

    if st.button("📤 Exportar laudo como .docx"):
        try:
            doc = Document(f"templates/{template_escolhido}")
            for par in doc.paragraphs:
                if "Paciente:" in par.text:
                    par.text = f"Paciente: {nome_paciente}             DN: {dn_paciente.strftime('%d/%m/%Y')}"
                elif "Data:" in par.text and "Médico solicitante:" in par.text:
                    par.text = f"Data: {data_exame.strftime('%d/%m/%Y')}       Médico solicitante: {medico_solicitante}"
            doc.add_paragraph("")
            doc.add_paragraph(st.session_state.texto_laudo)
            nome_arquivo = f"Laudo_{nome_paciente.replace(' ', '_')}.docx"
            doc.save(nome_arquivo)
            with open(nome_arquivo, "rb") as file:
                st.download_button(label="📥 Baixar laudo", data=file, file_name=nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.error(f"Erro ao exportar: {e}")

with col2:
    st.header("📄 Pré-visualização do Laudo")
    st.session_state.texto_laudo = st.text_area("Texto do laudo:", value=st.session_state.texto_laudo, height=600, label_visibility="collapsed")

streamlit_js_eval(js_expressions="""
(() => {
  const recognition = new (window.SpeechRecognition || window.webkitSpeechRecognition)();
  recognition.continuous = false;
  recognition.interimResults = false;
  recognition.lang = 'pt-BR';
  let timeoutId = null;
  recognition.onstart = () => {
    clearTimeout(timeoutId);
    timeoutId = setTimeout(() => recognition.stop(), 10000);
  };
  recognition.onresult = (event) => {
    const result = event.results[0][0].transcript;
    if (result.toLowerCase().startsWith("lili")) {
      window.parent.postMessage({ type: 'VOICE_RESULT', result }, '*');
    }
  };
  recognition.start();
})();
""", key="voz_coluna", trigger=False)

st.caption("Assistente Lili © 2025")
